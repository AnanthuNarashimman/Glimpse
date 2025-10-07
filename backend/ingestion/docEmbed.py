import os
import PyPDF2
import docx

from langchain_community.vectorstores import Chroma
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document


"""
This  module is used to process a document file (.pdf, .docx, .txt), chunks its text,
Create embeddings, and stores them in a specified ChromaDB collection.
"""


# Splits the processed text from documents into chunks using "RecursiveCharacterTextSplitter", embeds it and stores them in ChromaDB

def process_and_embed_document(file_path: str, collection_name: str, chroma_client, embedding_model):

    print(f"\n --- Ingesting document: {file_path} ---")

    full_text = _extract_text_from_file(file_path)

    if not full_text:
        print(f" Skipping '{file_path}' due to processing failure ot empty content.")
        return False
    
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
    text_chunks = text_splitter.split_text(full_text)
    print(f" - Split document into {len(text_chunks)} chunks.")

    documents = [
        Document(
            page_content=chunk,
            metadata = {"source": os.path.basename(file_path), "type": file_path.split('.')[-1]}
        ) for chunk in text_chunks
    ]

    try:
        print(f" - Storing {len(documents)} chunks in collection '{collection_name}'")

        Chroma.from_documents(
            documents=documents,
            embedding=embedding_model,
            collection_name=collection_name,
            client=chroma_client
        )

        print(f" - Successfully Stored embeddings in ChromaDB.")
        return True
    
    except Exception as e:
        print(f" - Error storing the documents in ChromaDB: {e}")




# Extraces text from the given document (.pdf, .docx, .txt)

def _extract_text_from_file(file_path: str) -> str:
    text = ""

    try:
        if file_path.endswith('.pdf'):
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text+=page_text
        
        elif file_path.endswith('.docx'):
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
        
        elif file_path.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()

        if not text.strip():
            print(f" - Warning: Extracted text is empty. It might be a scanned document.")
            return None
    
        print(f" - Successfully extracted {len(text)} cgaracters from '{os.path.basename(file_path)}'.")
        return text
    
    except Exception as e:
        print(f" - Error extracting text from the file {e}")
        return None
        
